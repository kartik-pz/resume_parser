import os
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
import uvicorn
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Security as FastAPISecurity, Request
from fastapi.security.api_key import APIKeyHeader
from fastapi.concurrency import run_in_threadpool
from pydantic import BaseModel, Field
from dotenv import load_dotenv

from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

from transformers import pipeline, AutoTokenizer
from optimum.intel import OVModelForTokenClassification # For loading OpenVINO model

# For text extraction
import docx
import pypdfium2 as pdfium

# --- Configuration ---
load_dotenv() # Load environment variables from .env file

API_KEY_NAME = "X-API-KEY"
API_KEY = os.getenv("API_KEY", "dev_key_123")
OV_MODEL_DIR_INT8 = Path("./openvino_model_int8")


# --- Global Variables / Application State ---
# To be populated by the load_model_and_tokenizer function at startup
app_state: Dict[str, Any] = {}

# --- Pydantic Models for API Response ---
class ParsedResumeResponse(BaseModel):
    PERSON: Optional[List[str]] = Field(default_factory=list, description="Extracted person names")
    Designation: Optional[List[str]] = Field(default_factory=list, description="Extracted job titles/designations")
    ORG: Optional[List[str]] = Field(default_factory=list, description="Extracted organizations/companies")
    GPE: Optional[List[str]] = Field(default_factory=list, description="Extracted geopolitical entities (locations)")
    Phone: Optional[List[str]] = Field(default_factory=list, description="Extracted phone numbers")
    Skills: Optional[List[str]] = Field(default_factory=list, description="Extracted skills")
    EducationDegree: Optional[List[str]] = Field(default_factory=list, description="Extracted education degrees")
    ExperianceYears: Optional[List[str]] = Field(default_factory=list, description="Extracted years of experience statements")
    DATE: Optional[List[str]] = Field(default_factory=list, description="Extracted dates")
    CARDINAL: Optional[List[str]] = Field(default_factory=list, description="Extracted cardinal numbers")


    class Config:
        json_schema_extra = {
            "example": {
                "PERSON": ["John Doe"],
                "Designation": ["Software Engineer", "Project Manager"],
                "ORG": ["Tech Solutions Inc.", "Innovate LLC"],
                "GPE": ["New York", "London"],
                "Phone": ["(123) 456-7890"],
                "Skills": ["Python", "FastAPI", "OpenVINO", "Problem Solving"],
                "EducationDegree": ["B.Sc. Computer Science", "MBA"],
                "ExperianceYears": ["5 years", "2+ years of experience"],
                "DATE": ["May 2020 - Present", "Jan 2018"],
                "CARDINAL": ["10", "Three"]
            }
        }

# --- Model Loading and Helper Functions ---
def load_model_and_tokenizer():
    """Loads the quantized OpenVINO model and tokenizer at application startup."""
    print(f"Loading quantized OpenVINO model from: {OV_MODEL_DIR_INT8}")
    if not OV_MODEL_DIR_INT8.exists():
        raise RuntimeError(f"OpenVINO model directory not found: {OV_MODEL_DIR_INT8}")

    try:
        tokenizer = AutoTokenizer.from_pretrained(OV_MODEL_DIR_INT8)
        ov_model_int8 = OVModelForTokenClassification.from_pretrained(
            OV_MODEL_DIR_INT8,
        )

        nlp_pipeline = pipeline(
            "token-classification",
            model=ov_model_int8,
            tokenizer=tokenizer,
            aggregation_strategy="simple"
        )
        
        print("OpenVINO model and pipeline loaded successfully.")
        return nlp_pipeline, tokenizer # Return tokenizer for post-processing
    except Exception as e:
        print(f"Error loading OpenVINO model or creating pipeline: {e}")
        raise RuntimeError(f"Could not load the model: {e}")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        pdf = pdfium.PdfDocument(file_bytes)
        text = ""
        for i in range(len(pdf)):
            page = pdf[i]
            textpage = page.get_textpage()
            text += textpage.get_text_range() + "\n" # Add newline between pages
            textpage.close()
            page.close()
        pdf.close()
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Could not process PDF file: {e}")


def extract_text_from_docx(file_path_or_stream: Any) -> str:
    try:
        doc = docx.Document(file_path_or_stream)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Could not process DOCX file: {e}")


def structure_resume_entities(pipeline_result: List[Dict[str, Any]], tokenizer_for_conversion) -> Dict[str, List[str]]:
    """
    Converts a flat list of entities from the HF pipeline into a structured dictionary.
    Handles non-B-I-O labels by grouping consecutive tokens of the same type.
    """
    structured_output: Dict[str, List[str]] = {}
    if not pipeline_result:
        return structured_output

    # The pipeline with aggregation_strategy="simple" should already group subwords.
    # Each item in pipeline_result should be a detected entity.
    for entity_item in pipeline_result:
        label = entity_item.get('entity_group') # 'entity_group' is typical for aggregated results
        if not label: # Fallback if 'entity_group' is not present
            label = entity_item.get('entity')
        
        text = entity_item.get('word', '').strip()

        if not label or label == "O" or not text: # Skip "O" labels or empty text
            continue

        # Correct potential typo in label from model training
        if label == "ExperianceYears":
            label = "ExperienceYears" # Standardizing the key

        if label not in structured_output:
            structured_output[label] = []
        
        # Basic deduplication for the same text under the same label
        if text not in structured_output[label]:
            structured_output[label].append(text)
            
    return structured_output


# --- FastAPI Application Setup ---
app = FastAPI(
    title="Optimized CV Parsing API",
    description="A scalable, high-performance API for parsing and extracting structured data from resumes using a quantized OpenVINO model.",
    version="1.0.0"
)

# Rate Limiter Setup
limiter = Limiter(key_func=get_remote_address, default_limits=["1000/hour", "50/minute"]) # Example limits
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# API Key Security
api_key_header_auth = APIKeyHeader(name=API_KEY_NAME, auto_error=False)

async def get_api_key(api_key_header: Optional[str] = FastAPISecurity(api_key_header_auth)):
    if not api_key_header or api_key_header != API_KEY:
        raise HTTPException(status_code=403, detail="Could not validate credentials: Invalid API Key")
    return api_key_header

# --- Application Lifespan Events (Startup/Shutdown) ---
@app.on_event("startup")
async def startup_event():
    """Load ML model and tokenizer when the application starts."""
    try:
        nlp_pipeline, tokenizer_instance = load_model_and_tokenizer()
        app_state["nlp_pipeline"] = nlp_pipeline
        app_state["tokenizer"] = tokenizer_instance # Store tokenizer if needed for post-processing
        print("Application startup: Model and tokenizer loaded.")
    except RuntimeError as e:
        print(f"FATAL: {e}")
        # Optionally, exit if model loading fails, or let it run and endpoints return error
        # For a critical service like this, exiting might be preferable
        # import sys
        # sys.exit(1)


# --- API Endpoints ---
@app.post("/parse-resume/",
            response_model=ParsedResumeResponse,
            tags=["CV Parsing"],
            summary="Parse a resume file and extract structured entities.")
@limiter.limit("30/minute") # Override default limits for this specific endpoint if needed
async def parse_resume_endpoint(
    request: Request,
    file: UploadFile = File(..., description="Resume file (PDF, DOCX, or TXT) to be parsed."),
    api_key: str = Depends(get_api_key) # Enforce API key authentication
):
    """
    Upload a resume file, parse its content using an optimized NLP model,
    and return the extracted structured data in JSON format.

    Supported file types: `application/pdf`, `application/vnd.openxmlformats-officedocument.wordprocessingml.document`, `text/plain`.
    """
    if "nlp_pipeline" not in app_state or "tokenizer" not in app_state:
        raise HTTPException(status_code=503, detail="Model not loaded. Service is temporarily unavailable.")

    content_type = file.content_type
    print(f"Received file: {file.filename}, Content-Type: {content_type}")

    text_content = ""
    try:
        file_bytes = await file.read()
        if not file_bytes:
            raise HTTPException(status_code=400, detail="Uploaded file is empty.")

        if content_type == "application/pdf":
            text_content = await run_in_threadpool(extract_text_from_pdf, file_bytes)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # python-docx needs a file-like object or path
            import io
            file_stream = io.BytesIO(file_bytes)
            text_content = await run_in_threadpool(extract_text_from_docx, file_stream)
        elif content_type == "text/plain":
            text_content = file_bytes.decode("utf-8", errors="ignore") # Be lenient with text encoding
        else:
            raise HTTPException(
                status_code=415,
                detail=f"Unsupported file type: {content_type}. Please upload PDF, DOCX, or TXT."
            )

        if not text_content.strip():
            raise HTTPException(status_code=400, detail="Extracted text from file is empty or malformed.")

    except ValueError as ve: # Specific error from text extraction
        raise HTTPException(status_code=422, detail=f"Error processing file content: {str(ve)}")
    except HTTPException: # Re-raise HTTPExceptions
        raise
    except Exception as e:
        print(f"Unhandled error during file processing: {e}")
        raise HTTPException(status_code=500, detail=f"Server error during file processing: {str(e)}")

    # Perform NLP inference in a thread pool as it's CPU-bound
    try:
        nlp_pipeline_instance = app_state["nlp_pipeline"]
        tokenizer_instance = app_state["tokenizer"]

        # The pipeline call is blocking
        pipeline_raw_result = await run_in_threadpool(nlp_pipeline_instance, text_content)
        
        # Structure the entities
        # The tokenizer_instance might not be strictly needed by structure_resume_entities
        # if pipeline_raw_result already contains clean words after aggregation.
        # Pass it anyway if your structuring logic uses it for subword conversion (though less likely now).
        structured_data = await run_in_threadpool(structure_resume_entities, pipeline_raw_result, tokenizer_instance)

    except Exception as e:
        print(f"Error during NLP inference or structuring: {e}")
        # Log the full traceback for debugging
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error during resume analysis: {str(e)}")

    # Ensure all keys from Pydantic model are present, even if with empty lists
    final_response = {key: structured_data.get(key, []) for key in ParsedResumeResponse.__fields__}
    
    return ParsedResumeResponse(**final_response)


# --- Main execution (for local development) ---
if __name__ == "__main__":
    if not API_KEY or API_KEY == "your_default_dev_api_key_in_code_if_not_in_env":
        print("WARNING: Running with a default or missing API_KEY. Set API_KEY environment variable for security.")
    if not OV_MODEL_DIR_INT8.exists():
        print(f"ERROR: Model directory {OV_MODEL_DIR_INT8} not found. The application might not start correctly or will fail on first request.")
    
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")